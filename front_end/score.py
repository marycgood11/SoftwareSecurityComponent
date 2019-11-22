#!/usr/bin/env python3
import datetime
import sys
import tkinter as tk
from tkinter import *
from front_end import scenarios
from docx import Document
from docx.shared import Inches

__author__ = "Mary Catherine Good"

"""
GLOBAL VARIABLES
"""
score_total = 0
score_percent = 0
security_rating = ""

"""
CONSTANTS
"""
SCENARIO1 = "Password Creation"
SCENARIO1_TIP = "When creating a new password, you should follow these best practice rules:\n\tInclude - 12+ characters, upper and lower case letters, numbers, special characters.\n\tNot Include - Personal information such as birthdays and names of pets and family, words that can be found in a dictionary, and any repeated sequences of letters or numbers.\nThe reason for this is because it is easy for hackers to find personal information about you such as your birthday, names of your family members, and the names of your pets.  Additionally, try not to include words that are found in a dictionary because hackers can use a simple algorithm to plug in words in an attempt to break your password.  Overall, the best passwords are unique passwords."
SCENARIO2 = "Password Usage"
SCENARIO2_TIP = "When it comes to best practices, you should never reuse passwords.  If a hacker gets a hold of one password than any website or account that uses that password has been compromised.  You should always have a separate password for important accounts such as bank accounts, email, and work applications.  If for some reason you have to use the same password for multiple account you should only use that password for non-important and frivolous accounts."
SCENARIO3 = "Physical Access"
SCENARIO3_TIP = "Whenever you leave your desk you should always lock your computer and take your badge with you.  Even if you are leaving for “just a second”, it only takes a hacker or rogue employee one second to take a badge.  Additionally, you should never hold doors open unless the employee has a security badge.  Too often, people with malicious intent are able to pose as if they work for the company and gain access by following others through the doors."
SCENARIO4 = "Separation of Privilege"
SCENARIO4_TIP = "Your access should be limited to the areas that you need to complete your job. If you were granted access to everything then it could become a security hazard. There is the potential that you could misuse data or expose information to shady people. Additionally, implementing this policy would make it hard for hackers to have access to high level data from exploiting a weakness from an asset lower in the chain of command. Check with your managers and department heads to see what access you will require."
SCENARIO5 = "Data Encryption"
SCENARIO5_TIP = "Sensitive data should always be encrypted without exception. This is especially true when sending information to others through less secure means such as email."
SCENARIO6 = "Phishing"
SCENARIO6_TIP = "First off, you should always check the email addresses of emails that you receive.  Hackers will try and send you emails that appear that they are from legit sources when in fact, they slightly altered the email address.  Additionally, if you find emails from outside sources in your inbox that you do not recognize, you should exercise the utmost caution. You should never open suspicious emails or open suspicious email links.  More often than not it is a hacker trying to put malicious software on your computer and/or gain information. If you think you received a phishing email, notify your cybersecurity group."
SCENARIO7 = "Bringing Data Home"
SCENARIO7_TIP = "You should never put work information of any sort onto a flash drive. Flash drives are small and easily lost. If a flash drive full of sensitive work data were to fall into the wrong hands, some serious damage could be wrought upon the company.  Additionally, it is a security hazard to use your own personal computer for work.  If you want to use your own device to work, you should take it to your company’s tech/security desk.  There they will ensure that your device is outfitted with all the proper security controls that are necessary."
SCENARIO8 = "VPN vs. Public Network"
SCENARIO8_TIP = "Public WiFi Networks are not the safest to connect to. Your data is vulnerable to hackers and could be intercepted.  When working outside of the office, it is best to connect to a VPN (Virtual Private Network).  VPN’s are safe, secure, encrypted, and hidden."
SCENARIO9 = "Data Management / Encryption"
SCENARIO9_TIP = "You should always encrypt sensitive and confidential information. Even though the company has taken measures to protect and prevent against cybersecurity attacks, attacks still do happen. If an attack was successful, then all of your teams or department sensitive and confidential information could be exposed. By encrypting data, you are creating another line of defense as well as ensuring your data stays safe."


class Score(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.controller = controller
        self.configure(background="#E6F1FF")

        # CREATE FRAMES
        title_frame = tk.Frame(self, background="#E6F1FF", width=550, pady=10)
        border_frame1 = tk.Frame(self, background="#5BA3FF", height=2, width=550)
        border_frame2 = tk.Frame(self, background="#006EFF", height=2, width=550)
        border_frame3 = tk.Frame(self, background="white", height=2, width=550)
        border_frame4 = tk.Frame(self, background="#006EFF", height=2, width=550)
        border_frame5 = tk.Frame(self, background="#5BA3FF", height=2, width=550)
        main_frame = tk.Frame(self, background="#E6F1FF", height=600, width=550)
        footer_frame = tk.Frame(self, background="#E6F1FF", height=200, width=550)

        # LAYOUT FRAMES
        title_frame.grid(row=0, column=0, sticky="nsew")
        border_frame1.grid(row=1, column=0, sticky="new")
        border_frame2.grid(row=2, column=0, sticky="new")
        border_frame3.grid(row=3, column=0, sticky="new")
        border_frame4.grid(row=4, column=0, sticky="new")
        border_frame5.grid(row=5, column=0, sticky="new")
        main_frame.grid(row=6, column=0)
        footer_frame.grid(row=7, column=0)

        # GET RESULTS WIDGETS
        security_awareness_label = tk.Label(title_frame, text="Security Awareness Training", font=("Noteworthy", 30), anchor="center", background="#E6F1FF", pady=15)
        congratulations_label = tk.Label(main_frame, text="Congratulations!", font=("Arial 30 bold"), wraplength=425, anchor="center", justify="left", background="#E6F1FF", pady=15)
        completion_label = tk.Label(main_frame, text="You have completed the SONY's Security Awareness Training.  \nClick Get Results to continue.", font=(20), wraplength=425, background="#E6F1FF", pady=10)
        place_holder_label = tk.Label(main_frame, background="#E6F1FF")
        results_button = tk.Button(main_frame, text="Get Results", font=(15), command=lambda: self.getResults(main_frame, footer_frame, congratulations_label, completion_label, place_holder_label, results_button), highlightbackground="#5BA3FF", pady=5, width=15)

        security_awareness_label.pack(side="top", fill="x")
        congratulations_label.pack()
        completion_label.pack()
        place_holder_label.pack()
        results_button.pack()


    def getResults(self, main_frame, footer_frame, congratulations_label, completion_label, place_holder_label, results_button):
        # DELETE OLD WIDGETS
        congratulations_label.destroy()
        completion_label.destroy()
        place_holder_label.destroy()
        results_button.destroy()

        # CALCULATE SCORE
        self.sumScore()
        # GET SECURITY RATING
        self.getSecurityRating()

        # SET UP NEW WIDGETS
        self.score_label = tk.Label(main_frame, text=score_percent, font=("Arial 20 bold"), wraplength=425, anchor="center", justify="left", background="#E6F1FF", pady=10)
        self.security_rating_label = tk.Label(main_frame, text=security_rating, font=("Arial 20 bold"), wraplength=425, anchor="center", justify="left", background="#E6F1FF", pady=10)
        self.place_holder_label = tk.Label(main_frame, background="#E6F1FF")
        self.suggestions_button = tk.Button(main_frame, text="Download Security Improvements", command=lambda: self.getSuggestions(), bg="#5BA3FF", highlightbackground="#5BA3FF", pady=5, width=25)
        self.submit_button = tk.Button(footer_frame, text="Submit", font=(15), command=lambda: self.endApplication(), bg="#5BA3FF", highlightbackground="#5BA3FF", state="disabled", pady=5, width=15)

        # PLACE WIDGETS ON THE FRAME
        self.score_label.pack()
        self.security_rating_label.pack()
        self.suggestions_button.pack()
        self.place_holder_label.pack()
        self.submit_button.pack(side="bottom")

    def sumScore(self):
        global score_total
        global score_percent

        for x in scenarios.score:
            score_total = score_total + scenarios.score[x]
        score_total = int((score_total / 27) * 100)
        score_percent = "Score: " + str(score_total) + "%"

    def getSecurityRating(self):
        global security_rating

        if score_total == 100:
            security_rating = "Security Rating: A+"
        elif score_total >= 90:
            security_rating = "Security Rating: A"
        elif score_total >= 80:
            security_rating = "Security Rating: B"
        elif score_total >= 70:
            security_rating = "Security Rating: C"
        elif score_total >= 60:
            security_rating = "Security Rating: D"
        else:
            security_rating = "Security Rating: FAILED"

        return security_rating

    def getSuggestions(self):
        global score_total
        global score_percent
        global security_rating
        self.submit_button["state"]="normal"

        # GET THE YEAR
        now = datetime.datetime.now()
        year = now.year

        # PRINT WORD DOCX
        document = Document()
        doc_name = "SecurityAwarenessTrainingTips" + str(year) + ".docx"
        doc_title = "Security Awareness Training " + str(year)

        document.add_heading(doc_title, 0)
        document.add_heading("Tips and Suggestions to Improve Security Awareness", level=1)

        document.add_heading("Quick Summary", level=2)
        document.add_paragraph(score_percent)
        document.add_paragraph(security_rating)

        document.add_heading("Security Tips", level=2)
        if score_total == 100:
            document.add_paragraph("You are a security pro. You are not only ensuring that you keep yourself safe but you are also helping keep SONY safe from cyberattacks. Congratulations!")
        else:
            # ADD SUGGESTIONS TABLE
            table = document.add_table(rows=1, cols=2)
            table.style = "Light Shading"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Topic"
            hdr_cells[1].text = "Tip"

            for x in scenarios.score:
                if scenarios.score[x] != 3:
                    self.print_tip(x, table)

        # SAVE WORD DOCUMENT
        document.save(doc_name)

    def print_tip(self, scenario_num, table):
        scenario = ""
        scenario_tip = ""

        # GET SCENARIO INFOMATION
        if scenario_num == 1:
            scenario = SCENARIO1
            scenario_tip = SCENARIO1_TIP
        elif scenario_num == 2:
            scenario = SCENARIO2
            scenario_tip = SCENARIO2_TIP
        elif scenario_num == 3:
            scenario = SCENARIO3
            scenario_tip = SCENARIO3_TIP
        elif scenario_num == 4:
            scenario = SCENARIO4
            scenario_tip = SCENARIO4_TIP
        elif scenario_num == 5:
            scenario = SCENARIO5
            scenario_tip = SCENARIO5_TIP
        elif scenario_num == 6:
            scenario = SCENARIO6
            scenario_tip = SCENARIO6_TIP
        elif scenario_num == 7:
            scenario = SCENARIO7
            scenario_tip = SCENARIO7_TIP
        elif scenario_num == 8:
            scenario = SCENARIO8
            scenario_tip = SCENARIO8_TIP
        else:
            scenario = SCENARIO9
            scenario_tip = SCENARIO9_TIP

        # ADD TO TABLE
        row = table.add_row()
        row.cells[0].text = scenario
        row.cells[1].text = scenario_tip

    def endApplication(self):
        # TERMINATE PROGRAM
        sys.exit()
